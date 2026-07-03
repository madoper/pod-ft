__anchor__ = "export"

import base64
import json
from datetime import UTC, datetime
from io import BytesIO
from typing import Any

from backend.apps.export.app.schemas.export import ExportRequest, ExportResponse


class ExportService:
    FORMAT_CONFIG: dict[str, dict[str, Any]] = {
        "json": {"content_type": "application/json", "ext": "json"},
        "docx": {
            "content_type": (
                "application/vnd.openxmlformats-officedocument"
                ".wordprocessingml.document"
            ),
            "ext": "docx",
        },
        "pdf": {"content_type": "application/pdf", "ext": "pdf"},
        "xlsx": {
            "content_type": (
                "application/vnd.openxmlformats-officedocument"
                ".spreadsheetml.sheet"
            ),
            "ext": "xlsx",
        },
    }

    async def export(self, payload: ExportRequest) -> ExportResponse:
        fmt = payload.format
        config = self.FORMAT_CONFIG.get(fmt, self.FORMAT_CONFIG["json"])

        if fmt == "json":
            data = self._build_json(payload)
        elif fmt == "docx":
            data = self._build_docx(payload)
        elif fmt == "pdf":
            data = self._build_pdf(payload)
        elif fmt == "xlsx":
            data = self._build_xlsx(payload)
        else:
            data = self._build_json(payload)

        timestamp = datetime.now(UTC).strftime("%Y%m%d_%H%M%S")
        safe_title = payload.title.replace(" ", "_").lower()[:40]
        filename = f"{safe_title}_{timestamp}.{config['ext']}"

        return ExportResponse(
            format=fmt,
            filename=filename,
            content_type=config["content_type"],
            data=data,
        )

    def _build_json(self, payload: ExportRequest) -> str:
        doc: dict[str, Any] = {
            "title": payload.title,
            "generated_at": datetime.now(UTC).isoformat(),
            "sections": [],
        }
        for s in payload.sections:
            doc["sections"].append({
                "title": s.title,
                "content": s.content,
                "citations": s.citations,
            })
        if payload.summary:
            doc["summary"] = payload.summary
        return json.dumps(doc, ensure_ascii=False, indent=2)

    def _build_docx(self, payload: ExportRequest) -> str:
        from docx import Document
        from docx.shared import Pt

        doc = Document()
        doc.add_heading(payload.title, 0)

        doc.add_paragraph(f"Generated: {datetime.now(UTC).strftime('%Y-%m-%d %H:%M UTC')}")
        doc.add_paragraph()

        for s in payload.sections:
            doc.add_heading(s.title, 1)
            doc.add_paragraph(s.content)
            if s.citations:
                for c in s.citations:
                    p = doc.add_paragraph(style="List Bullet")
                    run = p.add_run(f"Citation: {c}")
                    run.font.size = Pt(9)

        if payload.summary:
            doc.add_heading("Summary", 1)
            doc.add_paragraph(payload.summary)

        buf = BytesIO()
        doc.save(buf)
        return base64.b64encode(buf.getvalue()).decode()

    def _build_pdf(self, payload: ExportRequest) -> str:
        import os

        from fpdf import FPDF

        pdf = FPDF()
        pdf.add_page()

        font_paths = [
            "/usr/share/fonts/truetype/dejavu/DejaVuSans.ttf",
            "/usr/share/fonts/truetype/dejavu/DejaVuSans.ttf",
            "C:/Windows/Fonts/arial.ttf",
            "C:/Windows/Fonts/DejaVuSans.ttf",
        ]
        font_loaded = False
        for fp in font_paths:
            if os.path.exists(fp):
                pdf.add_font("CustomFont", "", fp)
                font_loaded = True
                break

        if not font_loaded:
            pdf.set_font("Helvetica", "", 10)

        def set_font(style="", size=10):
            if font_loaded:
                pdf.set_font("CustomFont", style, size)
            else:
                pdf.set_font("Helvetica", style, size)

        set_font(size=16)
        pdf.cell(0, 10, payload.title, new_x="LMARGIN", new_y="NEXT")
        set_font(size=8)
        gen_ts = datetime.now(UTC).strftime("%Y-%m-%d %H:%M UTC")
        pdf.cell(0, 10, f"Generated: {gen_ts}", new_x="LMARGIN", new_y="NEXT")
        pdf.ln(5)

        for s in payload.sections:
            set_font(size=12)
            pdf.cell(0, 10, s.title, new_x="LMARGIN", new_y="NEXT")
            set_font(size=10)
            pdf.multi_cell(0, 6, s.content)
            if s.citations:
                for c in s.citations:
                    set_font(size=8)
                    pdf.cell(0, 6, f"- Citation: {c}", new_x="LMARGIN", new_y="NEXT")
            pdf.ln(3)

        if payload.summary:
            set_font(size=12)
            pdf.cell(0, 10, "Summary", new_x="LMARGIN", new_y="NEXT")
            set_font(size=10)
            pdf.multi_cell(0, 6, payload.summary)

        buf = BytesIO()
        pdf.output(buf)
        return base64.b64encode(buf.getvalue()).decode()

    def _build_xlsx(self, payload: ExportRequest) -> str:
        from openpyxl import Workbook
        from openpyxl.styles import Font

        wb = Workbook()
        ws = wb.active
        ws.title = payload.title[:31]

        ws["A1"] = payload.title
        ws["A1"].font = Font(bold=True, size=14)
        ws["A2"] = f"Generated: {datetime.now(UTC).strftime('%Y-%m-%d %H:%M UTC')}"

        row = 4
        for s in payload.sections:
            ws.cell(row=row, column=1, value=s.title).font = Font(bold=True, size=11)
            row += 1
            ws.cell(row=row, column=1, value=s.content)
            row += 1
            if s.citations:
                for c in s.citations:
                    ws.cell(row=row, column=1, value=f"Citation: {c}")
                    row += 1
            row += 1

        if payload.summary:
            row += 1
            ws.cell(row=row, column=1, value="Summary").font = Font(bold=True)
            row += 1
            ws.cell(row=row, column=1, value=payload.summary)

        buf = BytesIO()
        wb.save(buf)
        return base64.b64encode(buf.getvalue()).decode()
